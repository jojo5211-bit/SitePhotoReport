from __future__ import annotations

import os
from dataclasses import dataclass
from pathlib import Path

from .store import ProjectStore


PDF_TEMPLATE_OPTIONS = (
    ("ink_saver", "黑白省墨版", "純黑白文字與線條，適合大量列印"),
    ("modern", "清爽藍綠版", "目前使用的清爽藍綠風格"),
    ("warm", "暖棕工程版", "暖棕與橘色調的工程紀錄風格"),
)

PDF_TEMPLATE_PALETTES = {
    "ink_saver": {
        "primary": "#000000",
        "secondary": "#222222",
        "accent": "#000000",
        "text": "#000000",
        "caption": "#222222",
        "footer": "#222222",
        "line": "#000000",
        "grid": "#777777",
        "panel": "#FFFFFF",
        "panel_alt": "#FFFFFF",
        "banner_bg": "#FFFFFF",
    },
    "modern": {
        "primary": "#143B50",
        "secondary": "#4B6470",
        "accent": "#29758A",
        "text": "#21343D",
        "caption": "#3F5660",
        "footer": "#6B7D84",
        "line": "#A9CFD8",
        "grid": "#C6DDE3",
        "panel": "#F2F8FA",
        "panel_alt": "#F8FBFC",
        "banner_bg": "#E5F2F4",
    },
    "warm": {
        "primary": "#553A2A",
        "secondary": "#80624B",
        "accent": "#B86D2C",
        "text": "#3F3027",
        "caption": "#6D5545",
        "footer": "#8A7567",
        "line": "#D7B38F",
        "grid": "#E6D2BE",
        "panel": "#FBF4EC",
        "panel_alt": "#FFF9F3",
        "banner_bg": "#F5E5D4",
    },
}


@dataclass
class ExportSettings:
    page_size: str = "A4"
    orientation: str = "portrait"
    photos_per_page: int = 4
    image_quality: int = 85
    pdf_template: str = "modern"
    show_cover: bool = True
    show_date: bool = True
    show_filename: bool = True
    show_location: bool = True
    show_caption: bool = True


def unique_output_path(directory: str | Path, base_name: str, extension: str) -> Path:
    """Return a non-overwriting output path, appending ``_1``, ``_2``... ."""
    directory = Path(directory)
    extension = extension if extension.startswith(".") else f".{extension}"
    candidate = directory / f"{base_name}{extension}"
    index = 1
    while candidate.exists():
        candidate = directory / f"{base_name}_{index}{extension}"
        index += 1
    return candidate


def project_info_lines(project) -> list[str]:
    """Build cover metadata without emitting placeholders for blank fields."""
    return [f"{label}：{str(value).strip()}" for label, value in project_info_pairs(project)]


def project_info_pairs(project) -> list[tuple[str, str]]:
    """Return only filled project metadata as label/value pairs."""
    fields = (
        ("工程名稱", project["name"]),
        ("工程編號", project["project_code"]),
        ("報告標題", project["report_title"]),
        ("工程地點", project["location"]),
        ("業主", project["client_name"]),
        ("承包商", project["contractor_name"]),
        ("承辦人", project["manager_name"]),
        ("公司名稱", project["company_name"]),
        ("開工日期", project["start_date"]),
        ("完工日期", project["end_date"]),
        ("工程說明", project["description"]),
    )
    return [(label, str(value).strip()) for label, value in fields if str(value or "").strip()]


def _font_path() -> Path | None:
    candidates = [
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "msjh.ttc",
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "mingliu.ttc",
        Path(__file__).resolve().parent / "assets" / "DejaVuSans.ttf",
    ]
    return next((path for path in candidates if path.exists()), None)


def _photo_path(store: ProjectStore, row) -> Path | None:
    if not row["placement_id"]:
        return None
    rendered = store.rendered_placement_path(row["placement_id"], max_side=1600)
    if rendered:
        return rendered
    relative = row["preview_rel_path"] or row["original_rel_path"]
    path = store.root / relative if relative else None
    return path if path and path.exists() else None


def _caption(row, settings: ExportSettings) -> str:
    pieces = []
    if settings.show_date and row["captured_at"]:
        pieces.append(str(row["captured_at"]))
    if settings.show_location and row["construction_location"]:
        pieces.append(str(row["construction_location"]))
    if settings.show_filename and row["original_filename"]:
        pieces.append(str(row["original_filename"]))
    if settings.show_caption and row["caption"]:
        pieces.append(str(row["caption"]))
    return "｜".join(pieces)


def export_pdf(store: ProjectStore, output: str | Path, settings: ExportSettings | None = None) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4, landscape, portrait
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import HRFlowable, Image, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from xml.sax.saxutils import escape

    settings = settings or ExportSettings()
    template_key = settings.pdf_template if settings.pdf_template in PDF_TEMPLATE_PALETTES else "modern"
    palette = PDF_TEMPLATE_PALETTES[template_key]

    def color(name: str):
        return colors.HexColor(palette[name])

    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    page_size = landscape(A4) if settings.orientation == "landscape" else portrait(A4)
    project = store.project()
    project_name = str(project["name"] or "工程照片紀錄")
    report_title = str(project["report_title"] or project_name)
    font_name = "Helvetica"
    font_path = _font_path()
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("SitePhotoCJK", str(font_path), subfontIndex=0))
            font_name = "SitePhotoCJK"
        except Exception:
            pass
    doc = SimpleDocTemplate(
        str(output), pagesize=page_size, rightMargin=14 * mm, leftMargin=14 * mm,
        topMargin=16 * mm, bottomMargin=17 * mm,
        title=report_title,
    )
    styles = getSampleStyleSheet()
    cover_kicker = ParagraphStyle("SiteCoverKicker", parent=styles["Normal"], fontName=font_name, fontSize=12, leading=16, textColor=color("accent"))
    cover_title = ParagraphStyle("SiteCoverTitle", parent=styles["Title"], fontName=font_name, alignment=TA_CENTER, fontSize=30, leading=38, textColor=color("primary"), spaceAfter=5 * mm)
    cover_subtitle = ParagraphStyle("SiteCoverSubtitle", parent=styles["Normal"], fontName=font_name, alignment=TA_CENTER, fontSize=14, leading=20, textColor=color("secondary"))
    cover_heading = ParagraphStyle("SiteCoverHeading", parent=styles["Heading2"], fontName=font_name, fontSize=18, leading=24, textColor=color("primary"), spaceBefore=4 * mm, spaceAfter=4 * mm)
    meta_label = ParagraphStyle("SiteMetaLabel", parent=styles["Normal"], fontName=font_name, fontSize=10.5, leading=15, textColor=color("secondary"))
    meta_value = ParagraphStyle("SiteMetaValue", parent=styles["Normal"], fontName=font_name, fontSize=12, leading=17, textColor=color("text"))
    heading_style = ParagraphStyle("SiteHeading", parent=styles["Heading2"], fontName=font_name, fontSize=20, leading=26, textColor=color("primary"), spaceBefore=2 * mm, spaceAfter=3 * mm)
    slot_style = ParagraphStyle("SiteSlot", parent=styles["Heading3"], fontName=font_name, fontSize=13.5, leading=18, textColor=color("accent"), spaceBefore=3 * mm, spaceAfter=2 * mm)
    body_style = ParagraphStyle("SiteBody", parent=styles["BodyText"], fontName=font_name, fontSize=10.5, leading=16, textColor=color("text"))
    caption_style = ParagraphStyle("SiteCaption", parent=body_style, fontSize=9.5, leading=13, textColor=color("caption"))
    note_style = ParagraphStyle("SiteNote", parent=body_style, alignment=TA_CENTER, fontSize=11, leading=16, textColor=color("secondary"))

    def on_page(canvas, document):
        canvas.saveState()
        canvas.setStrokeColor(color("grid"))
        canvas.setLineWidth(0.6)
        canvas.line(14 * mm, 11 * mm, page_size[0] - 14 * mm, 11 * mm)
        canvas.setFont(font_name, 8.5)
        canvas.setFillColor(color("footer"))
        canvas.drawString(14 * mm, 6 * mm, project_name)
        canvas.drawRightString(page_size[0] - 14 * mm, 6 * mm, f"第 {document.page} 頁")
        canvas.restoreState()

    story = []
    if settings.show_cover:
        pairs = project_info_pairs(project)
        description = next((value for label, value in pairs if label == "工程說明"), "")
        basic_pairs = [(label, value) for label, value in pairs if label != "工程說明"]
        metadata_rows = []
        for index in range(0, len(basic_pairs), 2):
            row = []
            for label, value in basic_pairs[index:index + 2]:
                row.extend([Paragraph(escape(label), meta_label), Paragraph(escape(value).replace("\n", "<br/>"), meta_value)])
            while len(row) < 4:
                row.extend(["", ""])
            metadata_rows.append(row)
        metadata_table = Table(metadata_rows or [["", "", "", ""]], colWidths=[25 * mm, 62 * mm, 25 * mm, 62 * mm], hAlign="LEFT")
        metadata_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), color("panel")),
            ("BOX", (0, 0), (-1, -1), 0.8, color("line")),
            ("INNERGRID", (0, 0), (-1, -1), 0.35, color("grid")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]))
        banner = Table([[Paragraph("SITE PHOTO REPORT", cover_kicker), Paragraph("工程照片整理與施工紀錄", cover_kicker)]], colWidths=[87 * mm, 87 * mm])
        banner_style = [
            ("BACKGROUND", (0, 0), (-1, -1), color("banner_bg")),
            ("BOX", (0, 0), (-1, -1), 0.8, color("line")),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ("ALIGN", (1, 0), (1, 0), "RIGHT"),
        ]
        if template_key == "warm":
            banner_style.append(("LINEBEFORE", (0, 0), (0, 0), 4, color("accent")))
        banner.setStyle(TableStyle(banner_style))
        story.extend([
            Spacer(1, 4 * mm), banner,
            Spacer(1, 18 * mm),
            Paragraph(escape(report_title), cover_title),
            Paragraph(escape(project_name), cover_subtitle),
            Spacer(1, 15 * mm),
            Paragraph("工程基本資料", cover_heading),
            metadata_table,
        ])
        if description:
            description_table = Table([[Paragraph("工程說明", meta_label), Paragraph(escape(description).replace("\n", "<br/>"), meta_value)]], colWidths=[25 * mm, 149 * mm], hAlign="LEFT")
            description_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), color("panel_alt")),
                ("BOX", (0, 0), (-1, -1), 0.8, color("line")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
            ]))
            story.extend([Spacer(1, 7 * mm), description_table])
        story.extend([
            Spacer(1, 25 * mm),
            HRFlowable(width="100%", thickness=1, color=color("line")),
            Spacer(1, 6 * mm),
            Paragraph("照片紀錄自下一頁開始", note_style),
            PageBreak(),
        ])
    rows = store.all_report_rows()
    grouped_photos = []
    current_key = None
    current_group = None
    for row in rows:
        path = _photo_path(store, row)
        if not path:
            continue
        try:
            source_width, source_height = ImageReader(str(path)).getSize()
        except (OSError, ValueError, TypeError):
            continue
        key = (row["section_id"], row["slot_id"])
        if key != current_key:
            current_key = key
            current_group = {
                "section_id": row["section_id"],
                "section_name": str(row["section_name"]),
                "slot_name": str(row["slot_name"]),
                "photos": [],
            }
            grouped_photos.append(current_group)
        current_group["photos"].append((row, path, source_width, source_height))

    has_photo = bool(grouped_photos)
    section_number = 0
    previous_section_id = None
    photo_page_started = False
    allowed_counts = (1, 2, 4, 6)
    photos_per_page = settings.photos_per_page if settings.photos_per_page in allowed_counts else 4

    def page_grid(count: int) -> tuple[int, int, float]:
        """Choose a grid while keeping partial groups compact.

        The image height is based on the selected page capacity, not only on
        the number of photos in this group.  That prevents a one-row group
        from expanding to almost a whole page and leaving no room for the
        next group.
        """
        if photos_per_page == 1:
            columns, capacity_rows = 1, 1.0
        elif photos_per_page == 2:
            columns, capacity_rows = 1, 2.0
        elif photos_per_page == 4:
            columns, capacity_rows = 2, 2.0
        else:
            columns, capacity_rows = 2, 3.0
        rows = max(1, (count + columns - 1) // columns)
        sizing_rows = capacity_rows + 0.25 if count < photos_per_page else capacity_rows

        content_height = page_size[1] - 33 * mm
        header_reserve = 30 * mm
        caption_reserve = 18 * mm
        row_gap = 2 * mm
        image_height = max(
            45 * mm,
            (content_height - header_reserve - caption_reserve * sizing_rows - row_gap * (sizing_rows - 1)) / sizing_rows,
        )
        return columns, rows, image_height

    def append_photo_page(group: dict, chunk: list[tuple], chunk_index: int) -> None:
        nonlocal photo_page_started, section_number, previous_section_id
        if photo_page_started:
            story.append(PageBreak())
        photo_page_started = True
        if group["section_id"] != previous_section_id:
            section_number += 1
            previous_section_id = group["section_id"]
            heading_flowables = [
                Paragraph(f"第 {section_number} 章　{escape(group['section_name'])}", heading_style),
                HRFlowable(width="100%", thickness=1, color=color("line"), spaceAfter=2 * mm),
            ]
        elif chunk_index > 0:
            heading_flowables = [
                Paragraph(f"第 {section_number} 章　{escape(group['section_name'])}（續）", slot_style),
            ]
        else:
            heading_flowables = []
        heading_flowables.append(Paragraph(f"- {escape(group['slot_name'])}", slot_style))

        columns, _rows, image_height = page_grid(len(chunk))
        usable_width = page_size[0] - 24 * mm
        cell_width = usable_width / columns
        image_width = max(cell_width - 8 * mm, 30 * mm)
        grid_rows = []
        for start in range(0, len(chunk), columns):
            row_cells = []
            for row, path, source_width, source_height in chunk[start:start + columns]:
                ratio = min(image_width / max(source_width, 1), image_height / max(source_height, 1))
                image = Image(str(path), width=source_width * ratio, height=source_height * ratio)
                image.hAlign = "CENTER"
                caption = escape(_caption(row, settings) or "照片")
                row_cells.append([image, Paragraph(caption, caption_style)])
            while len(row_cells) < columns:
                row_cells.append("")
            grid_rows.append(row_cells)
        for row_index, grid_row in enumerate(grid_rows):
            table = Table(
                [grid_row],
                colWidths=[cell_width] * columns,
                rowHeights=[image_height + 18 * mm],
                hAlign="LEFT",
            )
            table.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BACKGROUND", (0, 0), (-1, -1), color("panel_alt")),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("BOX", (0, 0), (-1, -1), 0.5, color("grid")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, color("grid")),
            ]))
            # Keep the item/field heading together with the first photo row.
            # Later rows remain independent so the rest of this same field
            # can continue naturally within its own page.
            if row_index == 0:
                story.append(KeepTogether([*heading_flowables, table]))
            else:
                story.append(table)

    for group in grouped_photos:
        photos = group["photos"]
        for start in range(0, len(photos), photos_per_page):
            chunk = photos[start:start + photos_per_page]
            append_photo_page(group, chunk, start // photos_per_page)
    if not has_photo:
        story.append(Paragraph("目前沒有可輸出的照片。", body_style))
    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    return output


def export_word(store: ProjectStore, output: str | Path, settings: ExportSettings | None = None) -> Path:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    settings = settings or ExportSettings()
    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    project = store.project()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(str(project["report_title"] or project["name"]))
    run.bold = True
    run.font.size = Pt(20)
    info = document.add_paragraph("\n".join(project_info_lines(project)))
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    rows = store.all_report_rows()
    current_section = None
    current_slot = None
    for row in rows:
        if row["section_id"] != current_section:
            if current_section is not None:
                document.add_page_break()
            current_section = row["section_id"]
            document.add_heading(str(row["section_name"]), level=1)
            current_slot = None
        if row["slot_id"] != current_slot:
            current_slot = row["slot_id"]
            slot_heading = document.add_paragraph(str(row["slot_name"]))
            slot_heading.runs[0].bold = True
        path = _photo_path(store, row)
        if not path:
            continue
        table = document.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(5.8))
        caption = cell.add_paragraph(_caption(row, settings))
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph()
    document.save(str(output))
    return output


def export_excel(store: ProjectStore, output: str | Path, settings: ExportSettings | None = None) -> Path:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Font, PatternFill

    settings = settings or ExportSettings()
    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "照片清冊"
    headers = ["編號", "工程項目", "欄位", "位置", "拍攝日期"]
    if settings.show_filename:
        headers.append("原始檔名")
    headers.extend(["照片說明", "照片"])
    sheet.append(headers)
    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="245B75")
        cell.alignment = Alignment(horizontal="center")
    number = 1
    rows = store.all_report_rows()
    for row in rows:
        if not row["placement_id"]:
            continue
        excel_row = sheet.max_row + 1
        values = [
            number, row["section_name"], row["slot_name"], row["construction_location"] or "",
            (row["captured_at"] if settings.show_date else "") or "",
        ]
        if settings.show_filename:
            values.append(row["original_filename"])
        values.extend([_caption(row, settings), ""])
        sheet.append(values)
        path = _photo_path(store, row)
        if path:
            try:
                picture = XLImage(str(path))
                picture.width = 120
                picture.height = 90
                sheet.add_image(picture, f"{('H' if settings.show_filename else 'G')}{excel_row}")
            except Exception:
                pass
        sheet.row_dimensions[excel_row].height = 72
        number += 1
    widths = {"A": 10, "B": 24, "C": 16, "D": 20, "E": 20, "F": 28, "G": 45, "H": 18}
    if not settings.show_filename:
        widths = {"A": 10, "B": 24, "C": 16, "D": 20, "E": 20, "F": 45, "G": 18}
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.paperSize = sheet.PAPERSIZE_A4
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0
    workbook.save(str(output))
    return output
